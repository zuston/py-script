#!coding:utf-8

import sys
import os
import xlrd
from docx import Document

class Transform(object):
    def __init__(self, excelFile, wordFile):
        self.excelFilePath = excelFile
        self.wordFilePath = wordFile

    def start(self):
        self.putWord()

    def loadExcel(self):
        headerList = []
        dataList = []
        if os.path.exists(self.excelFilePath) is False:
            print 'excel path error'
            return
        book = xlrd.open_workbook(self.excelFilePath, encoding_override='utf-8')
        sh = book.sheet_by_index(0)

        columns = sh.ncols
        for header in range(0, columns):
            headerList.append(header)

        for rx in range(1, sh.nrows):
            oneList = []
            for value in range(0, columns):
                oneList.append(unicode((sh.row(rx)[value].value)))
            dataList.append(oneList)

        return headerList, dataList

    def putWord(self):
        headerList, dataList = self.loadExcel()

        otherHeaderList = headerList[-8:]
        pHeaderList = headerList[0:-8]

        document = Document()
        document.add_heading(u"领导信息", 0)
        for line in dataList:
            otherDataList = line[-8:]
            pDataList = line[0:-8]
            print otherDataList[2]

            areaStr = ''
            for pline in pDataList:
                areaStr += pline.strip() + "  "
            document.add_heading(unicode(areaStr), level=3)
            p=document.add_paragraph("")
            pStr = ''
            pStr += otherDataList[0]
            pStr += '    '
            pStr += (otherDataList[2])
            p.add_run(unicode(pStr)).italic = True
            # document.add_paragraph(unicode(otherDataList[2]))
            document.add_paragraph(unicode(u"性别:  " + otherDataList[3]))
            document.add_paragraph(unicode(u'生年：  '+ otherDataList[4]))
            document.add_paragraph(unicode(u"籍贯:  " + otherDataList[5]))
            document.add_paragraph(unicode(u"学历:  " + otherDataList[6]))
            document.add_paragraph(unicode(u"简历:  " + otherDataList[7]))
            #
            document.add_paragraph(unicode(""))

        document.save(self.wordFilePath)

if __name__ == '__main__':
    tf = Transform("./data/societyTask/officerInfoExcel/county.xls","./data/societyTask/officerInfoWord/county.docx")
    tf.start()
    # tf = Transform("./data/societyTask/officerInfoExcel/province.xls", "./data/societyTask/officerInfoWord/province.docx")
    # tf.start()
    # tf = Transform("./data/societyTask/officerInfoExcel/viceProvince.xls", "./data/societyTask/officerInfoWord/viceProvince.docx")
    # tf.start()
