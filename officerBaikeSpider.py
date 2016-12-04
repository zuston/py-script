#!coding:utf-8

import sys
import time
import os
import xlrd
from pyTool.tool import const

sys.path.append("..")
from bs4 import BeautifulSoup as bs
from pyTool.tool.Http import *
from pyTool.tool.ZRedis import *
import xlwt
from docx import Document

class importExcel(object):
    def __init__(self):
        self.dataPath = './data/officerExcel/'

    def load(self, fileName):
        list = []
        dataSet = set()
        if os.path.exists(self.dataPath) and os.path.isfile(self.dataPath + fileName):
            excelName = self.dataPath + fileName
            tag_row = 1
            book = xlrd.open_workbook(excelName, encoding_override='utf-8')
            sh = book.sheet_by_index(0)

            for rx in range(tag_row, sh.nrows):
                province = str((sh.row(rx))[1].value)
                city = str(sh.row(rx)[3].value)
                area = str(sh.row(rx)[5].value)
                name = str(sh.row(rx)[7].value)
                # print province + ' ' + city + ' ' + area + ' ' + name
                if len(name) != 0 and name != 'N/A' and name != '' and name != ' ':
                    dataSet.add(province.strip() + ':' + city.strip() + area.strip() + ':' + name.strip())
            return dataSet
        else:
            print 'file path error'
            return dataSet



class exportBase(object):
    def __init__(self):
        pass

    def init(self,exportPath,exportName,data=None):
        pass

    def load(self,data):
        pass

    def export(self):
        pass


class exportExcel(exportBase):
    def __init__(self):
        self.data = None
        self.exportPath = None
        self.exportName = None

    def init(self,exportPath,exportName,data=None):
        self.exportPath = exportPath
        self.exportName = exportName
        if data!=None:
            self.data = data

    def load(self,data):
        self.data = data

    def export(self):
        if os.path.exists(self.exportPath+self.exportName):
            import xlrd
            from xlutils.copy import copy
            filename = self.exportPath+self.exportName
            rb = xlrd.open_workbook(filename,encoding_override='utf-8')
            sh = rb.sheet_by_index(0)
            wb = copy(rb)
            ws = wb.get_sheet(0)
            flag = int(((sh.row(0))[0].value))
            i = 0
            for line in self.data[:4]:
                ws.write(flag, i, unicode(line))
                i += 1
            ws.write(0,0,flag+1)
            wb.save(filename)
        else:
            excelWorkBook = xlwt.Workbook()
            sheet = excelWorkBook.add_sheet('sheet1')
            count = 0
            for line in self.data:
                sheet.write(1,count,unicode(line))
                count += 1
            sheet.write(0,0,2)
            # print self.exportPath+self.exportName
            # exit(1)
            excelWorkBook.save(unicode(self.exportPath+self.exportName))

class exportWord(exportBase):
    def __init__(self):
        self.data = None
        self.exportPath = None
        self.exportName = None

    def init(self,exportPath,exportName,data=None):
        self.exportPath = exportPath
        self.exportName = exportName
        if data!=None:
            self.data = data

    def load(self,data):
        self.data = data

    def export(self):
        fileName = self.exportPath+self.exportName
        try:
            if os.path.exists(fileName):
                document = Document(fileName)
            else:
                document = Document()
                document.add_heading(unicode("县长信息"), 0)

            document.add_heading(unicode(self.data[0]+"  "+self.data[1]), level=3)
            p = document.add_paragraph("")
            p.add_run(unicode(self.data[2])).italic = True
            document.add_paragraph(unicode(self.data[3]))
            for k,v in self.data[5].items():
                document.add_paragraph(unicode(str(k)+": "+str(v)))
            resumeP = document.add_paragraph(unicode("简历："))
            for list in self.data[4]:
                resumeP.add_run(unicode(str(list)+"\n"))

            document.save(fileName)
        except Exception:
            print 'fail to save to word'


'''
根据数据中的keyword和limitword来进行爬取

核心
'''
class spider(object):
    def __init__(self, string):
        self.requestUrl = 'http://baike.baidu.com/search/word?word='
        self.http = Http()
        self.secondUrl = 'http://baike.baidu.com'
        self.searchWord(string)

    def searchWord(self, string):
        splitList = string.split(':')
        self.keyword = splitList[-1]
        self.limitWord = splitList[:-1]

    def start(self):
        code, msg, res = self.http.open(self.requestUrl + self.keyword)
        if code == 200:
            try:
                result = self.analyze(res,self.limitWord)
                return result
            except Exception:
                return ERROR.PARSEERROR,'parse error'
        else:
            return ERROR.NETWORK,'network failure'

    def analyze(self, res, limitWord):
        self.soup = bs(res, 'html.parser')
        existFlag = self.soup.find('div', class_='no-result')
        if existFlag != None:
            tool.echo(self.keyword + '  未找到此人的关联信息')
            return ERROR.DATANONE, 'error'
        polysemant = self.soup.find('div', class_='polysemant-list')
        if polysemant == None:
            # 需要核实无异议的项目
            tool.echo(self.keyword + ' 无异议')
            code,res = ERROR.LEVEL1, self.parserData(res)
            for i in limitWord:
                if i in res.simpleInfo:
                    return code, res
            return ERROR.DATANONE,'error'
        else:
            max = [0, '']
            for line in polysemant.findAll('li'):
                content = line.text
                weight = 0
                for oneLimit in limitWord:
                    if oneLimit in content:
                        weight += 1
                if max[0] < weight:
                    max[0] = weight
                    max[1] = line
            if max[0] == 0:
                tool.echo(self.keyword + '  未找到此人的关联信息')
                return ERROR.DATANONE, 'error'

            info = max[1]
            if info.find('span', class_='selected') != None:
                tool.echo(self.keyword + '  已选择：' + info.text)
                return ERROR.LEVEL2, self.parserData(res)

            tool.echo(self.keyword + '  优先选择' + info.text)
            return ERROR.LEVEL3, self.parserData(self.polysemantAchive(info.a['href']))

    def polysemantAchive(self, url):
        code, msg, res = self.http.open(self.secondUrl + url)
        if code == 200:
            return res
        else:
            # raise Exception("network error")
            return 'network error'

    def parserData(self, res):
        keyword = self.keyword
        complexInfoList = []
        complexParamDict = {}
        simpleInfo = ''
        soup = bs(res, 'html.parser')
        # 获取keyword的剪短介绍和详细介绍
        infomation = soup.findAll('div', class_='para')
        if infomation != None:
            simpleInfo += infomation[0].text
            for info in infomation[1:]:
                complexInfoList.append(info.text)

        # 某些次要关键词
        complexParam = soup.find('div', class_='basic-info cmn-clearfix')
        if complexParam != None:
            for line in complexParam.findAll('dl', class_='basicInfo-block'):
                for one in line.findAll('dt', class_='name'):
                    complexParamDict[str(one.string).strip()] = str(one.find_next_sibling('dd', class_='value').string).strip()

        return self.instanceBean([keyword, simpleInfo, complexInfoList, complexParamDict])

    def instanceBean(self, list):
        oneBean = bean()
        oneBean.keyword = list[0]
        oneBean.simpleInfo = list[1]
        oneBean.complexInfoList = list[2]
        oneBean.complexParamDict = list[3]
        return oneBean


class bean(object):
    def __init__(self):
        self.keyword = ''
        self.simpleInfo = ''
        self.complexInfoList = None
        self.complexParamDict = None


class ERROR(object):
    const.NETWORK = -1
    const.DATANONE = -2
    const.PARSEERROR = -3
    const.LEVEL1 = 1
    const.LEVEL2 = 2
    const.LEVEL3 = 3

    # network error
    NETWORK = const.NETWORK
    # have none data
    DATANONE = const.DATANONE
    PARSEERROR = const.PARSEERROR
    # priority 1
    LEVEL1 = const.LEVEL1
    # priority 2
    LEVEL2 = const.LEVEL2
    # priority 3
    LEVEL3 = const.LEVEL3


    msg = {
        NETWORK:"network error",
        DATANONE:"have none data",
        LEVEL1:"confidence level-1",
        LEVEL2:"confidence level-2",
        LEVEL3:"confidence level-3",
        PARSEERROR:"parser error"
    }

class tool(object):
    @staticmethod
    def echo(string, flag=1):
        if flag:
            print string


class factory(object):
    def __init__(self):

        self.data = None

        self.importInstance = None
        self.importDict = {'EXCEL':importExcel(),}
        self.importChoice = None

        self.exportInstance = None
        self.exportDict = {'EXCEL':exportExcel(),'WORD':exportWord()}
        self.exportChoice = None

        self.exportPath = "./data/baikeOfficer/"
        self.exportFilename = "a.txt"

        self.redisConn = ZRedis()
        self.redisConn.setContainerName("baikeSet3")

    def importData(self,filename):
        producer = filename.split('.')[1]
        if producer=='xls' or producer=='xlsx':
            self.importChoice = 'EXCEL'
        if self.importDict.has_key(self.importChoice):
            producerInstance = self.importDict[self.importChoice]
            self.importInstance = producerInstance
            # 断点进行，如果redis中无数据，则进行导入数据，如果有，则不导入
            if os.path.exists(self.exportPath+self.exportFilename) is False:
                print 'flush the redis'
                self.redisConn.flush()
            if self.isRedisEmpty():
                print "++++import data from excel file"
                self.data = producerInstance.load(filename)
                self.pushRedis()
        else:
            raise Exception('fail to choose the file')

    def pushRedis(self):
        print '++++push data to redis'
        for one in self.data:
            self.redisConn.push(one)

    def popRedis(self):
        return self.redisConn.pop()

    def isRedisEmpty(self):
        return self.redisConn.isEmpty()

    def exportDataChoice(self,choice):
        self.exportChoice = choice

    def exportData(self,data,path,saveFilename):
        if self.exportDict.has_key(self.exportChoice):
            exportInstance = self.exportDict[self.exportChoice]
            exportInstance.init(path,saveFilename,data)
            exportInstance.export()
        else:
            raise Exception("chioce is error")

    def setExportFileName(self,filename):
        suffixDict = {"EXCEL":"xlsx","WORD":"docx"}
        self.exportFilename = filename+"."+suffixDict[self.exportChoice]

    def start(self):
        errorCount = 0
        # for i in self.data:
        #     dataList = i.split(':')[:-1]
        #     s = spider(i)
        #     data = s.start()
        #     code,res = data
        #     if code<0:
        #         res = [ERROR.msg[code],]
        #         errorCount += 1
        #     else:
        #         keyword = res.keyword
        #         simpleInfo = res.simpleInfo
        #         complexInfoList = res.complexInfoList if res.complexInfoList!=None else ["",]
        #         complexParamDict = res.complexParamDict if res.complexParamDict!=None else {"":"",}
        #         dataList.append(keyword)
        #         dataList.append(simpleInfo)
        #         dataList.append(complexInfoList)
        #         dataList.append(complexParamDict)
        #     #     todo
        #         self.exportData(dataList,self.exportPath,self.exportFilename)
        #         time.sleep(2)


        while self.isRedisEmpty()==False:
            print "rest request number:"+str(self.redisConn.len())
            i = self.popRedis()
            dataList = i.split(':')[:-1]
            s = spider(i)
            data = s.start()
            code,res = data
            if code<0:
                if code == ERROR.PARSEERROR:
                    print res
                res = [ERROR.msg[code],]
                errorCount += 1
            else:
                keyword = res.keyword
                simpleInfo = res.simpleInfo
                complexInfoList = res.complexInfoList if res.complexInfoList!=None else ["",]
                complexParamDict = res.complexParamDict if res.complexParamDict!=None else {"":"",}
                dataList.append(keyword)
                dataList.append(simpleInfo)
                dataList.append(complexInfoList)
                dataList.append(complexParamDict)
            #     todo
                self.exportData(dataList,self.exportPath,self.exportFilename)
                # time.sleep(1)

        print ""
        print "未捕获到的信息:"+str(errorCount)


if __name__ == '__main__':
    # b = base()
    # for i in b.readData('县委书记名单.xls'):
    #     # i为若干限定词和一个关键词组成的string
    #     s = spider(i)
    #     s.start()

    factory = factory()
    factory.exportDataChoice('WORD')
    factory.setExportFileName("partyWord4")
    factory.importData('县委书记名单.xls')
    factory.start()

